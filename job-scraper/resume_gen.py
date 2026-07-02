"""Resume generator — replicates the Claude /resume-generator skill via the OpenAI API,
but enforces the user's EXACT resume format and renders a matching .docx."""
import os
import json
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Known personal links (rendered as real hyperlinks in the contact line).
LINKS = {
    "LinkedIn": "https://www.linkedin.com/in/gouthamgunnala/",
    "GitHub": "https://github.com/Goutham-1012",
    "Portfolio": "https://goutham-1012.github.io/My-Portfolio/",
}
EMAIL = "gunnalagouthamreddy0@gmail.com"
PHONE = "913-406-5191"
HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# Every one of these companies must remain in the experience section (no role may be
# dropped during rewriting). Distinctive substrings of the real employers.
EXPECTED_COMPANIES = ["NextEra", "Fiserv", "Textron", "Lowe", "Siemens"]

OPENAI_URL = "https://api.openai.com/v1/chat/completions"
DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")  # never default to the pricey model


def _resolve_model(model=None):
    """Resolve the model at CALL time (env may not have loaded at import time)."""
    return model or os.environ.get("OPENAI_MODEL") or "gpt-4o-mini"


def _is_reasoning(model):
    """gpt-5* and o-series are reasoning models that reject a custom temperature."""
    m = (model or "").lower()
    return m.startswith("gpt-5") or m.startswith("o1") or m.startswith("o3") or m.startswith("o4")


def _chat_payload(model, messages, temperature):
    """Build a chat payload, omitting temperature for reasoning models (they only
    accept the default), so gpt-5-mini / gpt-5.4 / o-series don't 400."""
    payload = {"model": model, "messages": messages,
               "response_format": {"type": "json_object"}}
    if not _is_reasoning(model):
        payload["temperature"] = temperature
    return payload

FONT = "Century Gothic"
BASE_RESUME_PATH = os.path.join(os.path.dirname(__file__), "base_resume.txt")

SYSTEM_PROMPT = """ROLE
You are a top 10 percent resume strategist specializing in Data Scientist and AI/ML Engineer resumes. You understand ATS systems, recruiter screening behavior, and hiring manager expectations. Transform the resume so it strongly aligns with the target job description while staying authentic and credible.

Do NOT change: company names, locations, education, employment dates, total years of experience.
You MAY modify: professional summary, job titles (only when realistic), responsibilities, skills, technologies/tools, metrics, KPIs, dashboards.

STRATEGY
- Analyze the job description first: extract core technical skills, languages, frameworks, data/cloud platforms, ETL/pipeline tools, visualization tools, databases, domain knowledge, responsibilities, expected outcomes. Identify mandatory vs preferred skills and technical keywords.
- Every extracted technology/tool/platform/framework/language must appear in the Skills section and be reflected in experience bullets when realistic.
- Achieve greater than 90 percent ATS keyword alignment. Never copy sentences from the job description.
- Most recent role must include 100 percent of mandatory skills; preferred skills there too when realistic. Earlier roles reflect ~80-90 percent of relevant skills when plausible.
- Every experience/project bullet must include a technical action, the tools/technologies used, and a measurable outcome or operational impact.
- Maintain domain realism and a believable technology progression (earlier roles simpler, recent roles deeper).
- No EM dashes. No generic adjectives (dedicated, hardworking, adaptable, motivated, team player). No fluff. No keyword stuffing.

PRECEDENCE RULE (read carefully)
- For CONTENT QUALITY, follow ALL skill rules above EXACTLY: >90 percent ATS keyword alignment, mandatory skills present in the most recent role at 100 percent, every bullet contains a technical action + tools/technologies + a measurable outcome, skill distribution across roles, domain consistency, believable technology progression, no EM dashes, no generic adjectives, no copied JD sentences, post-generation keyword audit, and the final ATS analysis.
- For STRUCTURE AND FORMAT, follow the user's existing resume EXACTLY (this overrides the skill's "max 7 skills" and any other layout limit).

STRUCTURE REQUIREMENTS (from the user's resume — these override skill layout limits)
- "summary": 3 to 4 lines (NOT bullets), roughly 70% role/seniority POSITIONING and 30% tools. Open by positioning the candidate for the target role at their true level, with years of experience and 2 to 3 role-specific strengths plus production/business impact; weave in only the strongest tools naturally. Do NOT cram every JD keyword here (that belongs in skills), and never use generic openers (results-driven, detail-oriented, passionate, highly motivated, self-motivated).
- "skills": keep the categorized style of the original (one entry per category, "Category Name: item, item, item"), BUT FOCUS THE SECTION ON THE JOB DESCRIPTION. REMOVE categories and individual skills that are not relevant to the target role (for example, drop LLM / RAG / GenAI / knowledge-graph / finance categories when the JD is a computer-vision or robotics role). Keep only JD-relevant categories, reorder them so the most JD-critical appear first, and ADD new JD-specific categories and named tools/methods. Aim for roughly 8 to 12 tightly relevant categories. A focused, pruned skills section ranks better than a long unfocused one. EVERY skills entry MUST be a "Category Name: item, item, item" line that groups related skills; NEVER list a single keyword on its own line. Required JD keywords must be folded into the most appropriate category (for example "Robotics & Automation: Robotics, vision-based automation, robotic workcells").
  CATEGORY NAMES MUST BE CONVENTIONAL/STANDARD résumé groupings (e.g. "Programming & Query Languages", "Data Engineering & Pipelines", "Databases & Data Modeling", "Messaging & Streaming", "Cloud & Infrastructure", "MLOps & Deployment", "BI & Visualization"). Do NOT turn a JD requirement phrase into its own skill category — that reads as keyword stuffing. AVOID artificial categories like "Backpressure Management: event-driven messaging, idempotency, dead-letter handling" or "NoSQL Data Modeling: partition strategy, consistency tradeoffs, query cost optimization"; instead fold those terms into normal categories (e.g. "Messaging & Streaming: Kafka, idempotency, dead-letter queues, backpressure" and "Databases & Data Modeling: DynamoDB, partitioning, consistency tradeoffs"). Keep items as real tools/skills, not requirement sentences.
- "experience": keep every company, location, and date EXACTLY as in the original. EVERY role MUST have between 6 and 8 bullets (match the original resume's bullet counts; never fewer than 6 per role). When pruning or reframing, replace off-target bullets with JD-relevant ones rather than deleting them, so the count stays 6 to 8. APPLY THE SKILL'S ROLE POSITIONING to titles based on what the JD emphasizes: SQL / reporting / dashboards / KPIs => Data Analyst; ETL / Spark / Airflow pipelines => Data Engineer; data modeling / warehouse / transformations => Analytics Engineer; otherwise keep an AI/ML title. Reposition titles when the JD justifies it; never change the company, location, or dates.
  SENIORITY GUARD (critical): titles may match the JD's role FAMILY/specialty but must NEVER be under-leveled relative to the candidate's experience. The candidate has 6+ years, so the most recent role must stay mid/senior (e.g. "Senior AI Engineer", "Generative AI Engineer", "AI Software Engineer", "Lead/Staff …"). NEVER prepend "Junior", "Associate", "Intern", "Entry-level", "Trainee", or "Graduate" to any title, even if the job posting's own title says so. Keep each role's seniority at least as high as in the original résumé. The PROFESSIONAL SUMMARY opener must also describe the candidate at a level consistent with 6+ years — never "Junior" or "Associate".
- "projects": keep the Project Highlights section with project names and rewritten bullets.
- "education" and "certifications": keep exactly as provided.

HARD REQUIREMENTS (non-negotiable — verify before returning)
1. CREDIBILITY FIRST: the resume must read like a real senior engineer's career story, NOT a document reverse-engineered from the JD. It must be believable to a recruiter and defensible by the candidate in a technical interview. Strong ATS coverage matters, but when it conflicts with believability, choose believability.
1a. METRICS - believable MIX, not every bullet: roughly TWO-THIRDS of experience bullets carry a quantified result; the rest close on a concrete business or operational outcome with NO number. VARY the metric TYPE (latency, throughput, cost, accuracy, adoption, time saved, scale) and the phrasing. Use approximate scale where a precise figure would be hard to defend ("cut review effort by roughly a third", "about 40% faster in pilot testing"). Never end every bullet in "by NN%" and never fabricate precise domain record counts - both read as generated.
1b. BULLETS must be scannable: about 16 to 30 words, naming at most 3 to 4 tools each. Prefer showing the APPROACH and a design TRADEOFF ("used metadata filters and eval thresholds to hold latency within limits") over listing many tools. Open with a natural verb (Built, Designed, Led, Shipped, Scaled, Automated), NOT robotic openers ("Architected and expanded", "Operationalized", "Spearheaded", "Leveraged"). Project bullets about 14 to 26 words.
2. CRITICAL: It is NOT enough to add the job description's skills to the summary and the skills section. You MUST rewrite the PROFESSIONAL EXPERIENCE bullets so the candidate visibly DID work with the JD's core skills and technologies. Each mandatory JD skill must be demonstrated through a concrete accomplishment in an experience bullet (with tools + metric), not merely listed. If the JD targets a different specialty than the original resume (for example computer vision, robotics, or signal processing instead of LLMs), reframe the existing accomplishments so they credibly demonstrate the new specialty using the same companies, domains, and dates.
2a. ANCHOR BULLETS (highest priority): The two most recent roles must EACH lead with 2 to 3 strong "anchor" bullets that directly prove the JD's core responsibilities and DOMAIN — not just its tool names. Weave in the JD's domain workflows and process concepts as concrete work, e.g. for medical imaging: CT/MRI image processing, segmentation, image registration, 3D anatomy, image-guided navigation; for regulated/medical-device work: requirements, verification and validation (V&V), FDA / ISO 13485 / IEC 62304, risk controls; for safety-critical ML: failure mode and hazard analysis, edge-case identification, mitigation strategies; and full-lifecycle delivery: prototype algorithms -> production inference services with CI/CD, automated testing, and model monitoring. Each anchor bullet still needs specific tools + a quantified outcome. IMPORTANT: anchor bullets must obey rule 2c — where a domain workflow is a stretch for the candidate's real background, express it with transferable/adjacent, defensible wording (e.g. "[domain]-style data", "framework applicable to [domain]") rather than inventing precise first-hand domain claims or volumes.
2b. Any high-value JD keyword (especially the named programming languages like C++ and object-oriented design, plus domain/process terms) must appear in at least one EXPERIENCE bullet as real work, NOT only in the summary or skills — but only where it is credible for that company's domain. Prefer reframing an existing accomplishment over inventing an unrelated one.
2c. DEFENSIBILITY (ATS-optimized but believable — this is the target style): Maximize ATS keyword coverage WITHOUT fabricating hyper-specific claims the candidate could not defend in an interview. When the JD's domain differs from the candidate's real background, use realistic, transferable, HEDGED framing instead of asserting deep first-hand domain experience.
   PREFER wording like: "built a benefits-analytics dashboard framework", "worked with claims-like and healthcare-adjacent datasets", "applied [skill] to [domain]-style data", "designed pipelines applicable to [domain] reporting", "transferable to [domain] workflows".
   AVOID inventing precise foreign-domain volumes or artifacts that imply direct experience, e.g. "processed 50M claim rows", "5M employee benefit records", "12 provider feeds", "tested the Tuva Project connector at <employer>", or claiming regulated/clinical first-hand work the candidate never did.
   Keep the JD's exact keywords present for ATS, but for domain-specific terms the candidate has NOT directly used (niche tools, regulated programs, specific connectors), it is fine to surface them mainly in the SKILLS section and reference them in experience with adjacent/transferable wording. Every experience bullet must be something the candidate can comfortably explain and defend. Metrics must be plausible and tied to systems/efficiency/scale of work actually done, not to fabricated domain-specific record counts.
3. RECENT-ROLE WEIGHTING (do NOT make every company match the JD): tailor by recency so it reads as a real career, not five copies of the target job.
   - Most recent role: strongest match, about 80 to 90% of the JD's core skills, with 2 to 3 anchor bullets proving the core responsibilities.
   - Second role: strong but broader, about 60 to 75%.
   - Older roles: SUPPORTING foundation only, about 30 to 50% - keep them largely as the candidate's real prior work; do NOT retrofit the JD's specialty, tools, and domain onto every old company. Projects fill specific gaps.
   If an older role suddenly shows the same specialized JD stack as the recent role, it reads as fabricated.
4. Every core tool/technology named in the JD MUST appear in the "skills" section. Only the CORE skills need to be demonstrated in an experience bullet (in the recent roles) - do NOT force every tool into a bullet; most belong in skills only.
4a. When the JD lists ALTERNATIVES, include ALL named options verbatim, not just one. Examples: "Python or Java" -> list BOTH Python and Java; "AWS, Azure, or GCP" -> list all three. Do not silently drop an option just because the candidate favored another.
4b. When the JD mentions them, EXPLICITLY include these exact phrases in the skills section and reflect them in the summary/experience: agentic frameworks, SDK and ADK tools, rapid prototyping, solution hardening, secure development standards, responsible AI, data privacy, model governance, enterprise AI collaboration, and internal/business stakeholder engagement. Show stakeholder/enterprise collaboration and regulated-environment delivery inside experience bullets where credible.
5. Zero EM dashes (—) and zero EN dashes used as separators in any rewritten text. Zero generic adjectives.
6. Do not invent new companies, locations, or dates. Keep education and certifications verbatim.
7. Populate "mandatory_keywords" with the 8 to 12 most important technical skills/tools the JD requires (short tokens or 2-3 word phrases), and "preferred_keywords" with nice-to-have ones. These drive an automated audit, so every mandatory_keyword MUST appear in the skills section, and the CORE ones must be demonstrated in the recent roles (not necessarily every one, and not forced into older roles). Use the JD's exact terminology where it reads naturally.
8. PRUNE IRRELEVANCE: Remove or reframe any experience bullet, skill, or project that is not relevant to the target JD. Do not leave off-target content (e.g. credit scoring, financial fraud, knowledge graphs, diffusion image generation) in a resume aimed at a different specialty. Refocus everything on the JD.
9. DOMAIN CONSISTENCY: Keep each company's real industry domain. Never attribute finance/insurance/claims work to an aviation, energy, or retail employer. Reframe accomplishments within that company's actual industry.
10. NAME SPECIFIC TOOLS/METHODS that fit the JD's specialty when realistic, e.g. computer vision: YOLO, Faster R-CNN, Mask R-CNN, OpenCV, semantic segmentation; signal processing: FFT, wavelets, Kalman filters, sensor fusion, noise filtering; edge: ONNX Runtime, TensorRT, quantization. Spread these across skills and experience.
11. Write proof-based bullets in the candidate's own voice. Demonstrate the skill through a concrete accomplishment, not by restating the requirement.
11a. NO VERBATIM PLAGIARISM: Never copy the JD's distinctive multi-word phrases word-for-word or near-verbatim. A recruiter compares the resume to the JD — lifted phrases read as echoing requirements and destroy credibility. Example to AVOID: JD says "manage backpressure across async workloads and maintain clean service boundaries" and the bullet says "Managed backpressure ... ensuring clean service boundaries". Instead paraphrase into specifics, e.g. "Throttled a Kafka consumer with bounded queues and dead-letter retries so downstream services stayed decoupled during 5x traffic spikes." Reuse the JD's single keywords/tool names (for ATS), but not its phrases or sentence structure.
11b. SHOW THE "HOW", not just the "what": experience bullets should include the approach, mechanism, architecture, or decision that achieved the result — the technique used, not merely the outcome that mirrors the JD, and phrased differently from the JD.
12. TITLE HONESTY: keep titles truthful and defensible. You MAY swap the role FAMILY to match the JD's specialty (AI Engineer -> ML Engineer / Data Scientist / Analytics Engineer) and you must NOT under-level (no Junior/Associate/Intern for a 6+ year candidate). But do NOT INFLATE: never promote to Principal, Staff, Distinguished, Director, VP, Head, Chief, or Manager when the real title was Senior or an individual-contributor level. To signal higher SCOPE, keep the honest title and add a hyphen scope suffix, e.g. "Senior AI Engineer - Principal-level GenAI architecture and evaluation ownership". Keep each role's seniority consistent with the original resume.
13. KEYWORD DENSITY: keep it natural. Each important keyword should appear a LIMITED number of times across the whole document (roughly: summary once, skills once, one to two experience bullets, maybe one project) - NOT repeated ten or more times. Over-repeating a term (e.g. "MCP" fifteen times) looks forced.
14. HUMAN LANGUAGE: write like a strong engineer, not an ATS parser. NEVER join words with hyphens or underscores as fake tokens ("LLM-integration", "Prompt-engineering", "agentic-coding-workflows", "Access-Control", "Data-Security", "distributed_inference", "model_registry", "PhD_or_pursuing"). Write them naturally ("LLM integration", "prompt engineering", "access control", "distributed inference"). Keep genuine hyphenated terms (fine-tuning, real-time, multi-agent, end-to-end, role-based, retrieval-augmented).
15. DEPTH / DEFENSIBILITY GATE: for every bullet ask "could the candidate explain this architecture for 3 to 5 minutes in an interview, and is it believable at THIS company's real domain?" If not, simplify an over-specific stack to defensible concepts (e.g. "secure AI-to-tool integration using OAuth, audit logging, and sandboxed execution" instead of a precise proprietary stack) and use transferable framing for domains the candidate has not directly worked in.

OUTPUT
Return ONLY a JSON object (no markdown) with this exact shape:
{
  "name": "string",
  "contact": "string (single line: email | phone | LinkedIn | GitHub | Portfolio)",
  "summary": ["line 1", "line 2", "line 3", "line 4"],
  "skills": ["Category: items", "Category: items", ...],
  "experience": [
    {"title": "string", "dates": "string", "company_location": "string", "bullets": ["...", "..."]}
  ],
  "projects": [
    {"name": "string", "bullets": ["...", "..."]}
  ],
  "education": ["line", "line"],
  "certifications": ["line", "line"],
  "mandatory_keywords": ["skill", "tool", ...],
  "preferred_keywords": ["skill", "tool", ...],
  "analysis": "Plain-text ATS analysis: JD keyword breakdown, mandatory skills, preferred skills, before vs after keywords added, ATS alignment estimate, and how the resume was improved."
}
"""


# ---------------------------------------------------------------------------
# JD understanding: classify role family + seniority and extract structured needs,
# so the rewrite is tailored to THIS job instead of one-size-fits-all.
# ---------------------------------------------------------------------------
ROLE_FAMILIES = [
    "AI Engineer", "GenAI / LLM Engineer", "AI Platform Engineer", "MLOps Engineer",
    "Machine Learning Engineer", "Data Scientist", "Analytics Engineer",
    "Full-Stack AI Engineer", "Full-Stack Engineer", "Backend Engineer",
    "Computer Vision / Deep Learning Engineer", "Robotics / Autonomous Systems Engineer",
    "Data Engineer", "BI / Analytics Engineer", "Forward Deployed Engineer",
    "AI Architect", "Principal / Staff / Lead Engineer",
]

# Per-role-family emphasis (item 3 of the brief). Folded into the generation prompt.
ROLE_FAMILY_GUIDANCE = {
    "GenAI / LLM Engineer": (
        "RAG, embeddings, vector databases, prompt engineering, agentic workflows, "
        "tool/function calling, evaluation harnesses, hallucination reduction, structured "
        "outputs, model routing, observability, safety, and responsible AI."),
    "AI Platform Engineer": (
        "scalable infrastructure, Kubernetes, cloud platforms, CI/CD, Terraform, agent "
        "observability, AI gateways, security, identity, monitoring, platform patterns, "
        "developer enablement, and prototype-to-production workflows."),
    "MLOps Engineer": (
        "model lifecycle, experiment tracking, model registry, CI/CD, drift detection, "
        "monitoring, reproducibility, cloud/GPU resource optimization, containerization, "
        "and production reliability."),
    "Machine Learning Engineer": (
        "feature pipelines, training/evaluation workflows, inference, model serving, "
        "online/offline evaluation, performance, latency, data quality, and production ML "
        "systems."),
    "Data Scientist": (
        "experiment design, statistical modeling, feature engineering, A/B testing, "
        "forecasting/classification, evaluation rigor, and translating analysis into "
        "business decisions with measurable impact."),
    "Analytics Engineer": (
        "SQL, Python, dbt, Power BI/Tableau/Looker, data modeling, pipelines, data "
        "quality, stakeholder analysis, dashboards, and business-facing reporting."),
    "BI / Analytics Engineer": (
        "SQL, dbt, semantic/data modeling, Power BI/Tableau/Looker, governed metrics, "
        "data quality, and self-serve dashboards for business stakeholders."),
    "Full-Stack AI Engineer": (
        "React, TypeScript, Node.js/Python, APIs, WebSockets, PostgreSQL, Redis, queues, "
        "retries, real-time UX, user-facing AI workflows, observability, and product "
        "usability."),
    "Full-Stack Engineer": (
        "React/TypeScript front ends, Node.js/Python services, REST/GraphQL APIs, "
        "relational and cache stores, testing, and shipping reliable user-facing features."),
    "Backend Engineer": (
        "API design, microservices, databases, caching, queues, scalability, latency, "
        "reliability, testing, and clean service boundaries."),
    "Computer Vision / Deep Learning Engineer": (
        "PyTorch, TensorFlow, image/video pipelines, annotation, training datasets, "
        "real-time inference, TensorRT, ONNX, model optimization, sensor fusion, evaluation "
        "metrics, and deployment constraints."),
    "Robotics / Autonomous Systems Engineer": (
        "perception, sensor fusion, control, real-time systems, ROS, simulation, edge "
        "inference, safety/hazard analysis, and hardware-in-the-loop testing."),
    "Data Engineer": (
        "ingestion, ETL/ELT, Spark, Airflow/orchestration, warehousing, data modeling, "
        "streaming (Kafka), partitioning, data quality, and pipeline reliability."),
    "Forward Deployed Engineer": (
        "customer-facing delivery, rapid prototyping, integration, solution hardening, "
        "stakeholder communication, and turning ambiguous requirements into shipped systems."),
    "AI Architect": (
        "reference architectures, reusable patterns, design authority, cross-team "
        "influence, governance, security, scalability, and roadmap input."),
    "Principal / Staff / Lead Engineer": (
        "architecture ownership, reusable patterns and reference architectures, design "
        "authority, cross-team influence, mentoring senior/staff engineers, governance, "
        "and measurable enterprise impact."),
}
GENERIC_ROLE_GUIDANCE = (
    "the JD's named tools and core responsibilities as concrete production work, with "
    "system design, reliability, and measurable business impact appropriate to the role.")

# Seniority positioning (item 2). Detected seniority maps to one of these tiers.
SENIORITY_GUIDANCE = {
    "junior": (
        "hands-on delivery, technical execution, learning agility, implementation quality, "
        "and collaboration. Keep every claim grounded in execution, not ownership."),
    "mid": (
        "hands-on delivery and owning features end to end, solid implementation quality, "
        "collaboration, and emerging system thinking."),
    "senior": (
        "ownership of production systems, system design, reliability, cross-functional "
        "delivery, mentoring, and clear business impact."),
    "lead": (
        "architecture ownership, reusable patterns and reference architectures, design "
        "authority, cross-team influence, executive communication, governance, roadmap "
        "input, mentoring senior/staff engineers, and measurable enterprise impact."),
}


def _seniority_tier(seniority):
    """Map a detected seniority string to a guidance tier key."""
    s = (seniority or "").lower()
    if any(t in s for t in ("principal", "staff", "lead", "architect", "director", "head")):
        return "lead"
    if "senior" in s or "sr" in s:
        return "senior"
    if any(t in s for t in ("junior", "entry", "associate", "intern", "grad")):
        return "junior"
    return "mid"


def analyze_jd(jd, model=None):
    """One structured-extraction call that classifies the JD and pulls out the signals
    the rewrite needs (role family, seniority, must/preferred skills, domain, etc).
    Returns {} on any failure so generation degrades gracefully to the base prompt."""
    jd = (jd or "").strip()
    if not jd:
        return {}
    model = _resolve_model(model)
    system = ("You are an expert technical recruiter. Read the job description and return "
              "ONLY a JSON object describing it precisely. Infer sensibly from context.")
    user = (
        "JOB DESCRIPTION:\n" + jd[:6000] + "\n\n"
        "Return JSON with EXACTLY these keys:\n"
        '  "target_title": string,\n'
        '  "role_family": one of [' + "; ".join(ROLE_FAMILIES) + "; Other],\n"
        '  "seniority": one of [junior, mid, senior, lead, principal, staff, architect],\n'
        '  "years_experience": string (e.g. "5+"),\n'
        '  "must_have": array of short skill/tool tokens (the hard requirements),\n'
        '  "preferred": array of nice-to-have tokens,\n'
        '  "domain": string (industry/domain),\n'
        '  "leadership": one of [none, some, high],\n'
        '  "hands_on_coding": one of [low, medium, high],\n'
        '  "cloud": array of clouds/platforms named,\n'
        '  "business_facing": one of [low, medium, high],\n'
        '  "compliance": array of governance/security/regulatory needs,\n'
        '  "location": string (e.g. "remote", "onsite NYC", "hybrid"),\n'
        '  "success_signals": array of 3-6 outcomes the role is measured on.\n'
    )
    try:
        payload = _chat_payload(model, [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ], 0.2)
        resp = _openai_post(payload, timeout=180)
        if resp.status_code >= 400:
            return {}
        return json.loads(resp.json()["choices"][0]["message"]["content"]) or {}
    except (requests.RequestException, KeyError, ValueError):
        return {}


def _as_list(v):
    if isinstance(v, list):
        return [str(x).strip() for x in v if str(x).strip()]
    if isinstance(v, str) and v.strip():
        return [v.strip()]
    return []


def _build_system_prompt(analysis):
    """Append a JD-specific TAILORING BRIEF to the base prompt so tone, positioning,
    skills, and bullets all match the detected role family + seniority (items 1-3, 5)."""
    if not analysis:
        return SYSTEM_PROMPT
    fam = analysis.get("role_family") or "Other"
    role_emph = ROLE_FAMILY_GUIDANCE.get(fam, GENERIC_ROLE_GUIDANCE)
    tier = _seniority_tier(analysis.get("seniority"))
    sen_emph = SENIORITY_GUIDANCE[tier]
    must = ", ".join(_as_list(analysis.get("must_have"))[:18]) or "(infer from JD)"
    pref = ", ".join(_as_list(analysis.get("preferred"))[:14]) or "(none stated)"
    signals = "; ".join(_as_list(analysis.get("success_signals"))[:6]) or "(infer from JD)"
    cloud = ", ".join(_as_list(analysis.get("cloud"))) or "(none named)"
    compliance = ", ".join(_as_list(analysis.get("compliance"))) or "(none named)"
    brief = f"""

TAILORING BRIEF (derived from THIS job description - apply it across summary, skills, experience, and projects)
- Target role: {analysis.get('target_title', '(infer)')}
- Role family: {fam}
- Seniority (positioning, NOT a license to under-level the candidate): {analysis.get('seniority', 'mid')} ({tier} tier)
- Domain / industry: {analysis.get('domain', '(infer)')}
- Years expected: {analysis.get('years_experience', '(infer)')}
- Leadership expectation: {analysis.get('leadership', 'some')} | Hands-on coding: {analysis.get('hands_on_coding', 'high')} | Business-facing: {analysis.get('business_facing', 'medium')}
- Cloud/platform expectations: {cloud}
- Compliance / security / governance expectations: {compliance}
- Must-have skills (HIGHEST ATS priority - distribute naturally across summary + skills + experience, never as a bare keyword dump): {must}
- Preferred skills: {pref}
- Role-specific success signals to demonstrate with concrete, quantified accomplishments: {signals}

ROLE-FAMILY EMPHASIS: lead the most recent role and the skills section with {role_emph}
SENIORITY POSITIONING: shape the summary, bullets, and project emphasis around {sen_emph}
  (Positioning changes the TONE and CONTENT - ownership/architecture for senior+, execution/delivery for junior/mid - NOT just the title. Never under-level a 6+ year candidate, and never INFLATE the title either: keep the real level and, for a lead/principal-scoped JD, express the higher scope with a hyphen suffix like "Senior AI Engineer - Principal-level GenAI architecture ownership" rather than printing a Principal/Staff title.)

RECENT-ROLE WEIGHTING (credibility): make it a real career, not five copies of this job. Most recent role ~80-90% aligned; second role ~60-75%; older roles are SUPPORTING foundation (~30-50%) kept close to the candidate's real prior work - do not retrofit this JD's specialty/tools/domain onto every old company. Use 1-2 projects to fill gaps.

PROFESSIONAL SUMMARY (item 5): make it role-specific and seniority-specific - open with the target-role positioning at the candidate's true level, state years of experience if accurate, name the top 2-3 role-specific strengths, cite production/business impact, align to the domain, reflect the leadership/collaboration level, and weave the strongest tools in naturally. BANNED summary phrases: results-driven, detail-oriented, highly motivated, passionate, self-motivated, proven track record, team player, go-getter. Do not make every summary sound the same.

BULLET QUALITY (item 6): each bullet = action/ownership + technical approach (the HOW) + business/production impact + a credible measurable result. Do NOT pack many tools into one bullet, reuse the same metric style repeatedly, exaggerate, copy JD wording, or stuff keywords. BANNED weak openers: "worked on", "helped with", "responsible for", "assisted with", "involved in", "participated in". Prefer specific, defensible accomplishments a candidate can explain in an interview.

ATS WITHOUT STUFFING (item 8): preserve the must-have keywords verbatim for ATS, but DISTRIBUTE them naturally across summary, skills, experience, and projects. Never create standalone one-word keyword lines.
"""
    return SYSTEM_PROMPT + brief


def get_key():
    key = os.environ.get("OPENAI_API_KEY")
    if not key or key.startswith("sk-xxxx"):
        raise RuntimeError("OPENAI_API_KEY not set. Add it to your .env file.")
    return key


def _openai_post(payload, timeout=300, retries=2):
    """POST to OpenAI with a generous timeout and automatic retry on transient network
    timeouts / connection drops. Reasoning models (gpt-5-mini) plus the multi-pass audit
    loop can occasionally exceed a single request window, so one slow call shouldn't fail
    the whole resume. Raises the last exception only if every attempt fails."""
    headers = {"Authorization": f"Bearer {get_key()}", "Content-Type": "application/json"}
    last = None
    for _ in range(retries + 1):
        try:
            return requests.post(OPENAI_URL, headers=headers, json=payload, timeout=timeout)
        except (requests.Timeout, requests.ConnectionError) as e:
            last = e
    raise last


def load_base_resume():
    if os.path.exists(BASE_RESUME_PATH):
        with open(BASE_RESUME_PATH, encoding="utf-8") as f:
            return f.read()
    return ""


def generate_resume(resume_text, job_description, model=None, expected_companies=None):
    """Call OpenAI (JSON mode) and return the structured resume dict.

    `expected_companies` (per-profile) drives the keep-all-roles audit; when None it
    falls back to the default profile's EXPECTED_COMPANIES (backward compatible)."""
    resume_text = _normalize_text((resume_text or "").strip()) or load_base_resume()
    if not resume_text:
        raise ValueError("Original resume is required.")
    if not (job_description or "").strip():
        raise ValueError("Target job description is required.")
    model = _resolve_model(model)  # resolve here so .env is honored regardless of import order

    # 1) Understand the JD first, then tailor the whole prompt to it (role family,
    #    seniority, must/preferred skills, domain). Degrades to the base prompt on failure.
    analysis = analyze_jd(job_description, model)
    system_prompt = _build_system_prompt(analysis)

    user_msg = (
        "ORIGINAL RESUME (preserve this exact structure and formatting):\n"
        f"{resume_text}\n\n"
        "TARGET JOB DESCRIPTION:\n"
        f"{job_description.strip()}\n\n"
        "Rewrite the resume to align with the job description and the TAILORING BRIEF, "
        "then return the JSON object."
    )

    payload = _chat_payload(model, [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_msg},
    ], 0.4)

    resp = _openai_post(payload)
    if resp.status_code >= 400:
        raise RuntimeError(f"OpenAI HTTP {resp.status_code}: {resp.text[:300]}")
    content = resp.json()["choices"][0]["message"]["content"]
    data = json.loads(content)

    # 2) Automatic audit + repair loop (now JD/seniority/role-family aware, plus
    #    credibility, defensibility, vague-phrase, anti-plagiarism, and anti-stuffing
    #    checks). Keep the best draft seen, ranked by (fewest problems, highest ATS).
    #    Once the floor is cleared, chase only a few more passes of minor nits.
    best = data
    best_key = (float("inf"),)
    polished = 0
    for _ in range(12):
        problems = _audit(data, job_description, expected_companies, analysis)
        score = ats_score(data, job_description)
        key = (len(problems), -score)
        if key < best_key:
            best, best_key = data, key
        if not problems and score >= ATS_FLOOR:
            break
        if score >= ATS_FLOOR:
            polished += 1
            if polished >= 3:
                break
        data = _repair(data, job_description, problems, model, system_prompt)

    # 3) Final quality gate: if any CRITICAL issue survived the loop (missing role,
    #    under-leveled title, em dash, encoding artifact), spend one last repair on it.
    critical = _critical_issues(best, job_description, expected_companies, analysis)
    if critical:
        fixed = _repair(best, job_description, critical, model, system_prompt)
        if len(_audit(fixed, job_description, expected_companies, analysis)) <= len(
                _audit(best, job_description, expected_companies, analysis)):
            best = fixed

    # 4) Deterministic post-processing (no extra API cost): clean encoding/format
    #    artifacts, then enforce the 3-page length budget (trim oldest roles first).
    best = _normalize_resume(best)
    best = _trim_to_pages(best, max_pages=3)

    # 5) Internal multi-dimensional scoring (item 13) + the analysis, attached for
    #    transparency. render_docx / data_to_text ignore these private keys.
    best["ats_score"] = ats_score(best, job_description)
    best["_analysis"] = analysis
    best["_scores"] = internal_scores(best, job_description, analysis)
    best["_est_pages"] = round(_estimate_lines(best) / PAGE_LINES, 1)
    return best


ATS_FLOOR = 85  # keep repairing until here; lowered from 90 so the model stops
                # over-stuffing keywords just to hit the number (credibility > raw ATS)
METRIC_COVERAGE = 0.6   # only ~2/3 of bullets need a number; the rest close on outcomes
KEYWORD_MAX_MENTIONS = 5  # a single keyword repeated more than this reads as forced
MAX_TOOLS_PER_BULLET = 5  # too many named tools in one bullet is dense / not scannable
# Titles at/above these levels are inflations for a Senior/IC candidate — express the
# higher scope with a hyphen suffix instead of printing the inflated title.
_OVERLEVEL_TITLES = ("principal", "staff", "distinguished", "director", " vp", "vp ",
                     "head of", "chief", "manager", "architect")

# Concrete named technologies that, if present in the JD, must surface in the resume.
_RECENT_TECH = {"python", "sql", "airflow", "prefect", "llm", "llms", "agentic",
                "pytorch", "tensorflow", "spark", "kafka", "docker", "kubernetes"}
# Broader vocabulary required VERBATIM in the skills section when present in the JD.
# Includes languages/clouds that JDs often list as alternatives ("Python or Java",
# "AWS, Azure, or GCP") plus common AI-engineering practice phrases reviewers look for.
_SKILLS_TECH = _RECENT_TECH | {
    # languages
    "java", "javascript", "typescript", "c++", "scala", "go", "r",
    # clouds / infra
    "aws", "azure", "gcp", "terraform", "ci/cd", "jenkins", "github actions",
    # ml / genai
    "forecasting", "classification", "optimization", "statistics",
    "feature engineering", "orchestration", "pipelines", "rag", "fine-tuning",
    "langchain", "langgraph", "llamaindex", "prompt engineering", "agentic frameworks",
    "generative ai", "genai", "nlp", "computer vision", "vector database",
    "embeddings", "mlops", "mlflow", "sagemaker", "vertex ai", "bedrock",
    "onnx", "tensorrt", "sdk", "adk",
    # web / integration
    "react", "node.js", "fastapi", "rest", "graphql", "microservices", "api",
    # governance / practices reviewers flag
    "responsible ai", "model governance", "data privacy", "rapid prototyping",
    "solution hardening", "secure development", "enterprise", "stakeholder",
    "regulated", "compliance",
}


# Terms that should be LISTED in skills when the JD names them, but NOT forced to be
# "demonstrated" in an experience bullet (languages/clouds/infra/web/governance that
# would be fabrication to claim hands-on in every recent bullet).
_LISTED_ONLY = {
    "java", "javascript", "typescript", "c++", "scala", "go", "r",
    "aws", "azure", "gcp", "terraform", "ci/cd", "jenkins", "github actions",
    "react", "node.js", "rest", "graphql", "microservices", "api", "sdk", "adk",
    "responsible ai", "model governance", "data privacy", "rapid prototyping",
    "solution hardening", "secure development", "enterprise", "stakeholder",
    "regulated", "compliance",
}


def _jd_terms(jd, vocab):
    """Return vocab terms present in the JD, matched on alphanumeric boundaries so
    'go'/'r' don't match inside 'governance' and trailing punctuation (e.g. 'java.')
    is handled correctly."""
    import re
    low = jd.lower()
    found = set()
    for t in vocab:
        pattern = r"(?<![a-z0-9])" + re.escape(t) + r"(?![a-z0-9])"
        if re.search(pattern, low):
            found.add(t)
    return sorted(found)


_SYNONYMS = {
    "machine learning": ["ml", "deep learning", "ml models", "neural"],
    "computer vision": ["cv", "opencv", "image", "vision model", "visual"],
    "model deployment": ["deployment", "deployed", "serving", "inference service"],
    "model training": ["training", "trained", "fine-tuned", "fine-tuning", "retraining"],
    "data collection": ["data ingestion", "data acquisition", "ingestion", "collected data"],
    "signal processing": ["sensor data", "telemetry", "scada", "time series", "spectral"],
    "sensor calibration": ["calibration", "sensor configuration", "sensor"],
    "automated testing": ["test framework", "unit test", "integration test", "ci/cd", "test suite"],
    "edge devices": ["edge", "edge hardware", "edge deployment", "on-device"],
    "robotics": ["robotic", "automation", "autonomous"],
    "genai": ["generative ai", "gen ai", "gen-ai", "llm", "generative"],
    "generative ai": ["genai", "gen ai", "llm", "generative"],
    "llms": ["llm", "large language model", "large language models"],
    "llm": ["llms", "large language model", "large language models"],
    "prompt engineering": ["prompt design", "prompting", "prompt-engineering", "prompts"],
    "cloud platforms": ["aws", "azure", "gcp", "cloud"],
    "enterprise application integration": ["enterprise integration", "systems integration",
                                           "enterprise", "integration"],
    "rag": ["retrieval-augmented", "retrieval augmented", "retrieval-augmented generation"],
    "agentic frameworks": ["agentic", "agent framework", "multi-agent", "agents"],
    "agentic workflows": ["agentic", "agentic ai", "agent workflow", "multi-agent",
                          "multi agent", "agents", "workflow automation"],
    "function/tool calling": ["function calling", "tool calling", "function-calling",
                              "tool/function calling", "tool use", "tool invocation"],
    "function calling": ["function/tool calling", "tool calling", "function-calling", "tool use"],
    "tool calling": ["function calling", "function/tool calling", "tool use"],
    "llm evaluation": ["llm eval", "model evaluation", "evaluation harness",
                       "evaluation harnesses", "evals", "model performance metrics", "offline evaluation"],
    "guardrails": ["guardrail", "safety mechanism", "safety mechanisms", "content moderation", "safety"],
    "data privacy": ["privacy", "differential privacy", "pii", "data protection"],
    "vector databases": ["vector database", "vector db", "vector store", "embedding store",
                         "faiss", "pinecone", "weaviate", "chromadb", "milvus", "semantic search"],
    "vector database": ["vector databases", "vector store", "embedding store", "faiss", "pinecone"],
    "observability": ["monitoring", "tracing", "prometheus", "grafana", "telemetry"],
    "fine-tuning": ["fine tuning", "finetuning", "lora", "qlora", "peft"],
}

_STOP = {"and", "the", "of", "for", "with", "a", "an", "to", "in", "on"}


def _contains(text, keyword):
    """Semantic-ish presence test: direct substring, known synonym, or all
    significant (singularized) words of a multi-word keyword present."""
    k = keyword.lower().strip()
    if k in text:
        return True
    for syn in _SYNONYMS.get(k, []):
        if syn in text:
            return True
    words = [w.rstrip("s") for w in k.split() if w not in _STOP and len(w) > 2]
    if len(words) >= 2 and all(w in text for w in words):
        return True
    return False


# Core skills that are relevant to almost any engineering JD (never pruned).
_ALWAYS_RELEVANT = {"python", "sql", "java", "javascript", "c++", "git", "docker",
                    "kubernetes", "aws", "azure", "gcp", "cloud", "linux", "rest",
                    "api", "apis", "ci/cd", "bash", "scala", "go"}


def _irrelevant_skill_lines(d, jd, mand, pref):
    """Flag skill category lines whose items share NO token with the JD,
    the extracted keywords, or the always-relevant core set."""
    import re
    relevant = set(re.findall(r"[a-z][a-z0-9\+\#\.]{2,}", jd.lower()))
    relevant |= _ALWAYS_RELEVANT
    for k in mand + pref:
        relevant |= {w.rstrip("s") for w in k.lower().split()}
    bad = []
    for line in d.get("skills", []):
        tokens = {w.rstrip("s") for w in re.findall(r"[a-z][a-z0-9\+\#\.]{2,}", line.lower())}
        if tokens and not (tokens & relevant) and not any(t in relevant for t in tokens):
            bad.append(line)
    return bad


def _stuffed_skill_lines(d):
    """Flag keyword-stuffed skill entries where the items just repeat the category
    (e.g. 'Compliance: compliance', 'GenAI: genai', 'LLMs: llm, llms'), which read as
    fake categories rather than real skill groupings."""
    import re
    norm = lambda s: re.sub(r"[^a-z0-9]", "", s.lower())
    bad = []
    for line in d.get("skills", []):
        if ":" not in line:
            continue
        cat, items = line.split(":", 1)
        item_list = [i.strip() for i in items.split(",") if i.strip()]
        cat_n = norm(cat)
        items_n = {norm(i) for i in item_list}
        # stuffed if the items add nothing beyond the category name itself
        if not item_list:
            bad.append(line)
        elif items_n <= {cat_n} or all(i == cat_n or i in cat_n or cat_n in i for i in items_n):
            bad.append(line)
    return bad


_BANNED_SKILL_CATEGORIES = {
    "genai", "gen ai", "ai", "ml", "aiml", "ai/ml", "ai / ml", "misc", "other", "others",
    "keywords", "keyword", "skills", "technologies", "technology", "tools", "general",
    "additional", "miscellaneous", "various", "core", "core skills", "technical skills",
    "other skills", "additional skills",
}


def _umbrella_skill_lines(d):
    """Flag vague catch-all skill categories (GenAI:/AI:/Other:/Keywords:) — the model
    uses them to dump verbatim JD tokens, which reads as an unnatural keyword list."""
    import re
    norm = lambda s: re.sub(r"[^a-z0-9/ ]", "", s.lower()).strip()
    bad = []
    for line in d.get("skills", []):
        if ":" not in line:
            continue
        if norm(line.split(":", 1)[0]) in _BANNED_SKILL_CATEGORIES:
            bad.append(line)
    return bad


def _keyword_dump_lines(d):
    """Flag skill lines that dump lowercase JD *phrases* (concepts) instead of real,
    properly-cased tools. The tell: 2+ all-lowercase multi-word items (e.g. 'agentic
    workflows, function/tool calling, llm evaluation'), which reads as a keyword list."""
    bad = []
    for line in d.get("skills", []):
        if ":" not in line:
            continue
        items = [i.strip() for i in line.split(":", 1)[1].split(",") if i.strip()]
        phrase_dumps = [i for i in items if (" " in i or "/" in i) and i == i.lower()]
        if len(phrase_dumps) >= 2:
            bad.append(line)
    return bad


def _micro_category_lines(d):
    """Flag requirement-as-category over-fragmentation: several one-item skill categories
    (the model's other way of forcing a JD phrase in). Only fires when there are many."""
    singles = [l for l in d.get("skills", []) if ":" in l
               and len([i for i in l.split(":", 1)[1].split(",") if i.strip()]) <= 1]
    return singles if len(singles) >= 3 else []


# --- Credibility / human-language detectors (recruiter + hiring-manager pass) ---
_TOOL_VOCAB = {
    "python", "sql", "java", "javascript", "typescript", "c++", "scala",
    "pytorch", "tensorflow", "keras", "jax", "hugging face", "transformers", "langchain",
    "langgraph", "llamaindex", "autogen", "crewai", "semantic kernel", "openai", "anthropic",
    "claude", "bedrock", "azure openai", "vertex ai", "sagemaker", "mlflow", "kubeflow",
    "triton", "torchserve", "kserve", "ray", "docker", "kubernetes", "onnx", "tensorrt",
    "fastapi", "flask", "graphql", "grpc", "kafka", "spark", "airflow", "flink", "nifi",
    "dbt", "snowflake", "bigquery", "redshift", "databricks", "faiss", "pinecone", "weaviate",
    "chromadb", "milvus", "pgvector", "elasticsearch", "neo4j", "redis", "postgresql",
    "mongodb", "cassandra", "terraform", "helm", "pulumi", "jenkins", "github actions",
    "prometheus", "grafana", "opentelemetry", "streamlit", "gradio", "power bi", "tableau",
    "looker", "opencv", "yolo", "shap", "lime", "fairlearn", "lora", "qlora", "peft", "clip",
    "bert", "llama", "mistral", "gemini", "oauth", "entra id", "signalr", "mcp",
}

_LEGIT_HYPHENS = {
    "retrieval-augmented", "role-based", "chain-of-thought", "state-of-the-art",
    "object-oriented", "event-driven", "data-driven", "open-source", "end-to-end",
    "multi-agent", "real-time", "fine-tuning", "zero-shot", "few-shot", "cross-functional",
    "low-latency", "high-availability", "in-context", "on-premise", "well-architected",
    "back-end", "front-end", "full-stack", "human-in-the-loop", "time-series", "multi-modal",
    "multi-turn", "long-context", "in-house", "up-to-date", "self-serve",
}

# Second-halves that indicate a NOUN-noun join that should be a space ("LLM-integration",
# "Access-Control"). Compound ADJECTIVES (LLM-powered, SHAP-based, fine-tuned, GPU-accelerated)
# are legit and are NOT flagged.
_ATS_NOUN_SECONDS = {
    "integration", "engineering", "control", "security", "management", "optimization",
    "processing", "analysis", "architecture", "development", "deployment", "governance",
    "detection", "recognition", "automation", "orchestration", "modeling", "monitoring",
    "infrastructure", "operations", "intelligence", "workflow", "workflows",
}

_ROBOTIC_OPENERS = ("architected and", "operationaliz", "spearhead", "leverag", "utiliz",
                    "pioneer", "champion", "expanded an", "orchestrated and", "engineered and")


def _tool_heavy_bullets(bullets):
    """Bullets that name more than MAX_TOOLS_PER_BULLET distinct known tools (dense/stuffed)."""
    import re
    bad = []
    for b in bullets:
        low = b.lower()
        n = sum(1 for t in _TOOL_VOCAB
                if re.search(r"(?<![a-z0-9])" + re.escape(t) + r"(?![a-z0-9])", low))
        if n > MAX_TOOLS_PER_BULLET:
            bad.append(b)
    return bad


def _robotic_openers(bullets):
    """Bullets that start with AI/robotic-sounding verbs (item 14 human language)."""
    return [b for b in bullets if any(b.strip().lower().startswith(o) for o in _ROBOTIC_OPENERS)]


def _overused_keywords(d):
    """Mandatory keywords repeated more than KEYWORD_MAX_MENTIONS times across the resume."""
    import re
    content = json.dumps({k: v for k, v in d.items()
                          if k in ("summary", "skills", "experience", "projects")}).lower()
    out = []
    for k in {k.lower().strip() for k in d.get("mandatory_keywords", []) if k and len(k) > 1}:
        n = len(re.findall(r"(?<![a-z0-9])" + re.escape(k) + r"(?![a-z0-9])", content))
        if n > KEYWORD_MAX_MENTIONS:
            out.append((k, n))
    return sorted(out, key=lambda x: -x[1])[:8]


def _ats_token_lines(d):
    """Fake ATS tokens: words joined by underscores, or by hyphens where it should be a
    space (not a genuine hyphenated term). Human-language cleanup (item 14).
    Scans only visible content VALUES so JSON key names (company_location) don't match."""
    import re
    vals = (list(d.get("summary", [])) + list(d.get("skills", []))
            + list(d.get("education", [])) + list(d.get("certifications", [])))
    for e in d.get("experience", []):
        vals += [e.get("title", ""), e.get("company_location", "")] + e.get("bullets", [])
    for p in d.get("projects", []):
        vals += [p.get("name", "")] + p.get("bullets", [])
    blob = " ".join(x for x in vals if isinstance(x, str))
    found = set()
    for m in re.findall(r"\b[A-Za-z]{2,}(?:_[A-Za-z]{2,})+\b", blob):  # snake_case
        found.add(m)
    for m in re.findall(r"\b[A-Za-z]{2,}(?:-[A-Za-z]{2,}){2,}\b", blob):  # 3+ hyphen chains
        if m.lower() not in _LEGIT_HYPHENS:
            found.add(m)
    for m in re.finditer(r"\b([A-Za-z]{2,})-([A-Za-z]{2,})\b", blob):  # noun-noun joins only
        if m.group(0).lower() not in _LEGIT_HYPHENS and m.group(2).lower() in _ATS_NOUN_SECONDS:
            found.add(m.group(0))
    return sorted(found)[:12]


def _over_leveled_titles(d):
    """Titles inflated above the candidate's real Senior/IC level (item 12 title honesty).
    The honest scope suffix after a hyphen is ignored (only the base title is checked)."""
    bad = []
    for e in d.get("experience", []):
        t = e.get("title") or ""
        head = t.lower().split(" - ")[0].split(":")[0]
        if any(x.strip() in head for x in _OVERLEVEL_TITLES):
            bad.append(t)
    return bad


def _over_tailored_old_roles(exp_list, demo_required):
    """Older roles (3rd onward) that mirror too much of the JD's core stack (fabricated feel)."""
    if not demo_required or len(exp_list) <= 2:
        return []
    bad = []
    for e in exp_list[2:]:
        txt = json.dumps(e).lower()
        cov = sum(1 for k in demo_required if _contains(txt, k))
        if cov / len(demo_required) > 0.55:
            bad.append(f"{e.get('title', '?')} ({cov}/{len(demo_required)} core skills)")
    return bad


def ats_score(d, jd):
    """Keyword-coverage ATS estimate (0-100): how many JD-named skills appear in the
    resume, weighted toward demonstrating core specialty skills in experience.
    A proxy that correlates with external ATS tools, not an identical number."""
    mand = [k.lower().strip() for k in d.get("mandatory_keywords", [])
            if k and len(k.split()) <= 3]
    required = set(mand) | set(_jd_terms(jd, _SKILLS_TECH))
    if not required:
        return 100
    # Score against the actual resume content only (exclude the keyword/analysis fields
    # so they can't trivially satisfy coverage).
    content = {k: v for k, v in d.items()
               if k in ("name", "contact", "summary", "skills",
                        "experience", "projects", "education", "certifications")}
    full = json.dumps(content).lower()
    present = sum(1 for t in required if t in full or _contains(full, t))
    coverage = present / len(required)

    demo = [k for k in mand if k not in _LISTED_ONLY]
    exp = json.dumps(d.get("experience", [])).lower()
    demo_cov = (sum(1 for k in demo if _contains(exp, k)) / len(demo)) if demo else 1.0

    return max(0, min(100, round(100 * (0.75 * coverage + 0.25 * demo_cov))))


def _audit(d, jd, expected_companies=None, analysis=None):
    """Return concrete, fixable problems (with the exact offending content).

    `expected_companies` is the per-profile employer list; defaults to the built-in
    EXPECTED_COMPANIES. An empty list skips the keep-all-roles check entirely.
    `analysis` (from analyze_jd) drives the seniority/role-family alignment checks."""
    import re
    issues = []

    # No original role may be dropped during rewriting.
    companies = EXPECTED_COMPANIES if expected_companies is None else expected_companies
    exp_blob = json.dumps(d.get("experience", [])).lower()
    missing_roles = [c for c in companies if c.lower() not in exp_blob]
    if missing_roles:
        issues.append("These required roles are MISSING and must be restored with their real "
                      "company, location, dates, and 6-8 bullets (never drop a role): "
                      + ", ".join(missing_roles))

    # Metrics: aim for a believable MIX (~2/3 quantified), not every bullet a number.
    exp_bullets = [b for e in d.get("experience", []) for b in e.get("bullets", [])]
    quantified = [b for b in exp_bullets if re.search(r"\d", b)]
    if exp_bullets and len(quantified) / len(exp_bullets) < METRIC_COVERAGE:
        need = max(1, int(round(METRIC_COVERAGE * len(exp_bullets))) - len(quantified))
        issues.append(f"Too few bullets carry a measurable result (only {len(quantified)} of "
                      f"{len(exp_bullets)}). Add credible, VARIED metrics to about {need} more so "
                      "roughly two-thirds are quantified - vary the metric type (latency, cost, "
                      "throughput, accuracy, adoption, time saved) and use approximate scale where "
                      "a precise number would be hard to defend. Leave the rest closing on a "
                      "business/operational outcome without a number.")

    # Recent two roles carry the weight (6-8 bullets); older roles are lighter (>=4).
    too_few = []
    for i, e in enumerate(d.get("experience", [])):
        n = len(e.get("bullets", []))
        floor = 6 if i < 2 else 4
        if n < floor:
            too_few.append(f"{e.get('title', '?')} ({n} bullets, needs >={floor})")
    if too_few:
        issues.append("These roles have too few bullets. Recent roles need 6 to 8; older roles "
                      "need at least 4 (add credible ones, reframing real work): "
                      + ", ".join(too_few))

    # Scannability: not too thin, not too dense, not tool-stuffed.
    short = [b for b in exp_bullets if len(b.split()) < 12]
    if short:
        issues.append("These experience bullets are too thin (<12 words). Expand to about 16-30 "
                      "words with the approach and the outcome (no filler):\n   - "
                      + "\n   - ".join(short))
    longb = [b for b in exp_bullets if len(b.split()) > 34]
    if longb:
        issues.append("These bullets are too long/dense to scan. Tighten to about 16-30 words, "
                      "keeping the strongest point:\n   - " + "\n   - ".join(longb[:6]))
    tool_heavy = _tool_heavy_bullets(exp_bullets)
    if tool_heavy:
        issues.append(f"These bullets name too many tools (more than {MAX_TOOLS_PER_BULLET}); it "
                      "reads as keyword stuffing and hurts scannability. Keep the 2-4 most "
                      "relevant tools and show the approach/tradeoff instead:\n   - "
                      + "\n   - ".join(tool_heavy[:6]))

    skills_txt = json.dumps(d.get("skills", [])).lower()
    exp_txt = json.dumps(d.get("experience", [])).lower()
    recent_txt = json.dumps(d["experience"][0]).lower() if d.get("experience") else ""

    # Use the model's extracted JD keywords (short tokens/phrases) as the source
    # of truth; fall back to a fixed tech vocabulary if absent.
    mand = [k.lower().strip() for k in d.get("mandatory_keywords", [])
            if k and len(k.split()) <= 3]
    pref = [k.lower().strip() for k in d.get("preferred_keywords", [])
            if k and len(k.split()) <= 3]
    if not mand:
        mand = _jd_terms(jd, _SKILLS_TECH)

    # Every concrete technology / practice phrase NAMED in the JD must appear in the
    # skills section verbatim — this catches alternatives the model drops (e.g. it keeps
    # Python but not Java, AWS but not GCP) and reviewer-flagged phrases (agentic
    # frameworks, rapid prototyping, etc.). Union with the model's mandatory keywords.
    jd_required = sorted(set(mand) | set(_jd_terms(jd, _SKILLS_TECH)))

    # A JD skill counts as present if it appears verbatim OR via a known synonym/concept
    # already in the skills section. Concept-aware (not strict verbatim) so the model can
    # phrase awkward tokens naturally (e.g. "Agentic AI", "Function Calling") inside real
    # categories instead of dumping lowercase tokens or inventing one-off categories.
    miss_skills = sorted({k for k in jd_required if not _contains(skills_txt, k)}
                         | {k for k in pref if not _contains(skills_txt, k)})
    if miss_skills:
        issues.append("Represent these JD skills in the skills section (verbatim where it reads "
                      "naturally, otherwise as the standard tool/term for the concept), folded "
                      "into an appropriate conventional category: " + ", ".join(miss_skills))

    # Skill entries must be "Category: items" lines, not bare keywords.
    bare = [s for s in d.get("skills", []) if ":" not in s or not s.split(":", 1)[1].strip()]
    if bare:
        issues.append("These skills are bare keywords, not 'Category: items' lines. Fold each "
                      "into an appropriate category grouped with related skills (e.g. put "
                      "'Robotics' inside a 'Robotics & Automation: ...' line); never list a "
                      "keyword on its own line: " + ", ".join(bare))

    # Under-leveled titles (candidate has 6+ years) and a "Junior/Associate" summary opener.
    _BANNED_LEVEL = ("junior", "associate", "intern", "entry-level", "entry level",
                     "trainee", "graduate")
    bad_titles = [e.get("title", "") for e in d.get("experience", [])
                  if any(b in (e.get("title") or "").lower() for b in _BANNED_LEVEL)]
    summary0 = (d.get("summary") or [""])[0].lower()
    if bad_titles or any(b in summary0 for b in _BANNED_LEVEL):
        issues.append("Remove the under-leveled qualifier (Junior/Associate/Intern/Entry-level/"
                      "Trainee/Graduate) from titles and the summary — the candidate has 6+ years. "
                      "Use mid/senior titles (e.g. 'Senior AI Engineer', 'Generative AI Engineer', "
                      "'AI Software Engineer'). Offending: " + ", ".join(bad_titles or ["summary opener"]))

    # Keyword-stuffed lines where the category just repeats its own keyword.
    stuffed = _stuffed_skill_lines(d)
    if stuffed:
        issues.append("These skill entries are keyword-stuffing (the category just repeats its own "
                      "keyword). DELETE each fake category and FOLD the keyword as an item into an "
                      "appropriate REAL category (e.g. GenAI -> 'Generative AI & Agentic Workflows', "
                      "React/TypeScript -> 'Programming & Scripting', compliance -> 'Security & "
                      "Responsible AI') so the keyword stays present without a one-word category:\n   - "
                      + "\n   - ".join(stuffed))

    # Vague catch-all categories used to dump verbatim JD tokens (keyword-list smell).
    umbrella = _umbrella_skill_lines(d)
    if umbrella:
        issues.append("These skill lines use a vague catch-all category (GenAI/AI/ML/Other/"
                      "Keywords/Skills) that reads as a keyword dump. DELETE the umbrella "
                      "category and FOLD each item, properly capitalized, into the most specific "
                      "REAL category (e.g. move 'prompt engineering, LLM evaluation, function/tool "
                      "calling' into 'Generative AI & Agentic Workflows'; 'LLM' belongs in 'Large "
                      "Language Models & Deep Learning'). Keep the keywords present, just not as a "
                      "standalone dump:\n   - " + "\n   - ".join(umbrella))

    # Lowercase JD-phrase dumps hiding under a real-sounding category name.
    dumps = [l for l in _keyword_dump_lines(d) if l not in umbrella]
    if dumps:
        issues.append("These skill lines dump lowercase JD phrases as a list. Capitalize each "
                      "item properly (e.g. 'function/tool calling' -> 'Function/Tool Calling', "
                      "'agentic workflows' -> 'Agentic Workflows', 'llm evaluation' -> 'LLM "
                      "Evaluation') and FOLD them into the most specific REAL category (Generative "
                      "AI, Security & Responsible AI, etc.); delete the standalone dump line. Keep "
                      "every keyword present, just properly cased and grouped:\n   - "
                      + "\n   - ".join(dumps))

    # Requirement-as-category over-fragmentation (many one-item categories).
    micro = _micro_category_lines(d)
    if micro:
        issues.append("You have too many one-item skill categories (turning JD requirements into "
                      "their own categories reads as stuffing). Consolidate into roughly 8-12 "
                      "CONVENTIONAL categories (Programming, Generative AI & Agentic Workflows, "
                      "MLOps & Deployment, Vector Databases & Retrieval, Cloud & Infrastructure, "
                      "Security & Responsible AI, etc.) and fold each lone item into the right "
                      "standard group:\n   - " + "\n   - ".join(micro))

    # Prune skill categories that have no relevance to the JD at all.
    irrelevant = _irrelevant_skill_lines(d, jd, mand, pref)
    if irrelevant:
        issues.append("REMOVE these skill categories entirely — they are not relevant to this "
                      "JD and dilute the match:\n   - " + "\n   - ".join(irrelevant))

    # Demonstration follows RECENT-ROLE WEIGHTING (credibility): prove the CORE skills in
    # the recent role(s); do NOT force them into older roles, and flag older roles that are
    # over-tailored so the resume reads as a real career, not five copies of the JD.
    demo_required = [k for k in mand if k not in _LISTED_ONLY]
    exp_list = d.get("experience", [])
    recent_two = json.dumps(exp_list[:2]).lower()
    miss_recent = [k for k in demo_required if not _contains(recent_two, k)]
    if demo_required and len(miss_recent) > max(1, int(len(demo_required) * 0.25)):
        issues.append("The two most recent roles should demonstrate more of the JD's CORE skills "
                      "as concrete accomplishments (recent role strongest, ~80-90%): "
                      + ", ".join(miss_recent))

    over_old = _over_tailored_old_roles(exp_list, demo_required)
    if over_old:
        issues.append("These OLDER roles are over-tailored - they mirror the same specialized JD "
                      "stack as the recent role, which reads as fabricated. Make them the "
                      "candidate's real prior foundation (~30-50% overlap): keep broadly relevant "
                      "work, move JD-specific proof to the recent roles, and do NOT retrofit this "
                      "JD's niche tools/domain onto them: " + ", ".join(over_old))

    if "—" in json.dumps(d):
        issues.append("Remove all EM dashes.")
    full = json.dumps(d).lower()
    bad_adj = [w for w in ("dedicated", "hardworking", "adaptable", "motivated", "team player") if w in full]
    if bad_adj:
        issues.append("Remove generic adjectives: " + ", ".join(bad_adj))

    # --- Bullet quality: vague/weak openers (item 6) ---
    all_bullets = [b for e in d.get("experience", []) for b in e.get("bullets", [])]
    vague = [b for b in all_bullets if any(re.search(r"\b" + re.escape(p) + r"\b", b.lower())
                                           for p in VAGUE_PHRASES)]
    if vague:
        issues.append("Rewrite these bullets to remove weak/vague phrasing (worked on, helped "
                      "with, responsible for, assisted with, involved in, participated in) and "
                      "lead with a strong ownership verb + the technical approach + a metric:\n"
                      "   - " + "\n   - ".join(vague[:8]))

    # --- Professional summary: generic filler phrases (item 5) ---
    summary_txt = " ".join(d.get("summary", [])).lower()
    gen_sum = [p for p in GENERIC_SUMMARY_PHRASES if p in summary_txt]
    if gen_sum:
        issues.append("Remove these generic summary phrases and replace with role-specific, "
                      "seniority-specific positioning (top strengths + production/business "
                      "impact + domain): " + ", ".join(gen_sum))

    # --- Anti-plagiarism: bullets that copy a long JD phrase verbatim (item 6, 11a) ---
    lifted = _verbatim_overlaps(all_bullets, jd)
    if lifted:
        issues.append("These bullets copy the JD's wording too closely (a recruiter compares "
                      "side by side). Paraphrase into a specific, defensible accomplishment that "
                      "keeps only the single keywords, not the JD's phrases:\n   - "
                      + "\n   - ".join(lifted[:6]))

    # --- Defensibility: repeated metric style reads as fabricated (item 6, 7) ---
    monotony = _metric_monotony(all_bullets)
    if monotony:
        issues.append(monotony)

    # --- Encoding / formatting artifacts that must never reach the DOCX (item 11) ---
    artifacts = _encoding_artifacts(d)
    if artifacts:
        issues.append("Fix these text/encoding artifacts (broken apostrophes like \"Lowe20s\" -> "
                      "\"Lowe's\", stray symbols, doubled spaces): " + ", ".join(artifacts[:8]))

    # --- Title honesty (item 12): no under-leveling, no inflation; scope via a suffix ---
    over_lvl = _over_leveled_titles(d)
    if over_lvl:
        issues.append("These titles are INFLATED for a Senior/individual-contributor candidate "
                      "(Principal/Staff/Director/VP/Head/Chief/Manager/Architect). Revert to the "
                      "honest level; if the JD is lead/principal-scoped, express the scope with a "
                      "hyphen suffix instead, e.g. 'Senior AI Engineer - Principal-level GenAI "
                      "architecture ownership'. Never print an inflated standalone title: "
                      + ", ".join(over_lvl))
    if analysis and _seniority_tier(analysis.get("seniority")) == "lead":
        recent_title = (d.get("experience") or [{}])[0].get("title", "")
        if "-" not in recent_title:
            issues.append("The JD is lead/principal-scoped: add a scope suffix to the most recent "
                          "title (keeping the honest level), e.g. 'Senior AI Engineer - "
                          "Principal-level GenAI architecture and evaluation ownership', and lean "
                          "the recent bullets toward architecture ownership, reusable patterns, "
                          "cross-team influence, and mentoring.")

    # --- Keyword density (item 13): a term repeated too many times reads as forced ---
    overused = _overused_keywords(d)
    if overused:
        issues.append("These keywords are over-repeated and look forced. Reduce each to a few "
                      "natural mentions (summary once, skills once, 1-2 bullets, maybe one "
                      "project): " + ", ".join(f"{k} x{n}" for k, n in overused))

    # --- Human language (item 14): robotic verb openers (fake ATS tokens are cleaned
    # deterministically in _normalize_text, so they don't need a repair pass here) ---
    robotic = _robotic_openers(all_bullets)
    if robotic:
        issues.append("These bullets open with robotic/AI-sounding verbs; start with a plain, "
                      "strong verb (Built, Designed, Led, Shipped, Scaled, Automated) instead:\n"
                      "   - " + "\n   - ".join(robotic[:6]))

    if ats_score(d, jd) < ATS_FLOOR:
        issues.append(f"ATS keyword coverage is below the {ATS_FLOOR}% target. Ensure every "
                      "JD-named skill appears verbatim in the skills section and the core "
                      "specialty skills are demonstrated in the experience bullets.")
    return issues


def _repair(d, jd, problems, model, system_prompt=SYSTEM_PROMPT):
    """Send the draft back to the model with a targeted fix list."""
    msg = (
        "Here is a draft resume JSON. Make MINIMAL edits to fix ONLY the listed problems. "
        "Keep EVERYTHING ELSE byte-for-byte identical: same JSON shape, same companies, "
        "locations, dates, education, and certifications. Keep titles identical UNLESS a problem "
        "says to fix a title: for under-leveled, restore mid/senior wording; for INFLATED "
        "(Principal/Staff/Director/etc.), revert to the honest level and, only if asked, add a "
        "hyphen scope suffix ('Senior AI Engineer - Principal-level ... ownership'). Never change "
        "the company/dates. "
        "If a problem says to REMOVE/FOLD a skill category, delete that fake category line and merge its "
        "keyword as an item into a real category. NEVER drop an experience role; every original company must remain "
        "with its real location and dates. If a problem says a role is missing, restore it with "
        "6-8 bullets. If a role has too few bullets, ADD credible JD-relevant bullets (16-30 words) "
        "so it has 6 to 8. If a problem says an OLDER role is over-tailored, make it read as the "
        "candidate's real prior foundation: keep only broadly relevant work and move JD-specific "
        "proof to the two most recent roles - do NOT retrofit the JD's niche tools/domain onto old "
        "companies. "
        "Otherwise do not drop any bullet or skill not in the problem list, and do not add filler. "
        "When a CORE skill must be DEMONSTRATED, rewrite a bullet in the RECENT roles so the "
        "candidate genuinely did that work within the company's real domain (keep it defensible in "
        "an interview; use transferable framing for domains not directly worked in). "
        "Keep bullets scannable (16-30 words, at most 3-4 tools each), write natural English (no "
        "hyphen/underscore ATS tokens, no robotic openers like 'Architected and expanded'), and "
        "keep metrics a believable MIX (about two-thirds quantified with varied metric types; the "
        "rest close on an outcome without a number). Return the corrected JSON only.\n\n"
        "PROBLEMS TO FIX:\n" + "\n".join(problems) +
        "\n\nJOB DESCRIPTION:\n" + jd.strip() +
        "\n\nDRAFT JSON:\n" + json.dumps(d)
    )
    payload = _chat_payload(model, [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": msg},
    ], 0.3)
    try:
        resp = _openai_post(payload)
    except requests.RequestException:
        return d  # transient network failure — keep the draft
    if resp.status_code >= 400:
        return d  # keep the draft if repair call fails
    try:
        return json.loads(resp.json()["choices"][0]["message"]["content"])
    except (KeyError, ValueError):
        return d


# ---------------------------------------------------------------------------
# Quality helpers: bullet-quality / credibility / defensibility / formatting
# checks, internal scoring, length control, and text normalization.
# ---------------------------------------------------------------------------
VAGUE_PHRASES = ("worked on", "helped with", "responsible for", "responsible of",
                 "assisted with", "involved in", "participated in", "tasked with",
                 "duties included", "in charge of")
GENERIC_SUMMARY_PHRASES = ("results-driven", "results driven", "detail-oriented",
                 "detail oriented", "highly motivated", "passionate", "self-motivated",
                 "self motivated", "go-getter", "proven track record", "team player",
                 "hard-working", "hardworking", "dynamic professional",
                 "seasoned professional")


def _verbatim_overlaps(bullets, jd, n=6):
    """Bullets sharing an exact n-word run with the JD (near-verbatim plagiarism)."""
    import re
    toks = lambda s: re.findall(r"[a-z0-9\+\#]+", s.lower())
    jd_tokens = toks(jd)
    jd_grams = {" ".join(jd_tokens[i:i + n]) for i in range(len(jd_tokens) - n + 1)}
    if not jd_grams:
        return []
    bad = []
    for b in bullets:
        bt = toks(b)
        if any(" ".join(bt[i:i + n]) in jd_grams for i in range(len(bt) - n + 1)):
            bad.append(b)
    return bad


def _metric_monotony(bullets):
    """Flag templated bullets: one opener verb or one exact metric reused too often."""
    import re, collections
    if len(bullets) < 6:
        return ""
    msgs = []
    openers = collections.Counter((b.split() or [""])[0].lower().rstrip(",.;:") for b in bullets)
    top_word, cnt = openers.most_common(1)[0]
    if top_word and cnt >= max(4, int(0.4 * len(bullets))):
        msgs.append(f'{cnt} of {len(bullets)} bullets start with "{top_word}"')
    by_metrics = collections.Counter(re.findall(r"by \d+ ?%", " ".join(bullets).lower()))
    for phrase, c in by_metrics.items():
        if c >= 4:
            msgs.append(f'the metric "{phrase}" is reused {c} times')
    if msgs:
        return ("Vary the phrasing and metric style so bullets do not read as templated ("
                + "; ".join(msgs) + "). Use different strong verbs and a mix of metric types "
                "(latency, throughput, cost, accuracy, time saved, scale, adoption).")
    return ""


def _encoding_artifacts(d):
    """Detect text/encoding artifacts that must not reach the DOCX (item 11)."""
    import re
    blob = json.dumps(d, ensure_ascii=False)
    bad = []
    bad += sorted(set(re.findall(r"[A-Za-z]{2,}20s\b", blob)))  # Lowe20s, Lee20s
    for ch in ("â€", "Ã", "�"):
        if ch in blob:
            bad.append("mojibake")
            break
    if any(c in blob for c in ("\x91", "\x92", "\x93", "\x94", "\x95", "\x96", "\x97")):
        bad.append("Windows-1252 control chars")
    return bad


def _critical_issues(d, jd, expected_companies=None, analysis=None):
    """The subset of audit issues that must not survive to the final document."""
    issues = _audit(d, jd, expected_companies, analysis)
    keys = ("missing", "under-leveled", "inflated", "em dash",
            "below the", "encoding artifacts")
    return [i for i in issues if any(k in i.lower() for k in keys)]


def internal_scores(d, jd, analysis=None):
    """Multi-dimensional 0-100 self-scores (item 13). Informational; not shown unless
    the UI asks. Heuristic proxies, computed without extra API calls."""
    import re
    content = {k: v for k, v in d.items() if not k.startswith("_")}
    txt = json.dumps(content).lower()
    bullets = [b for e in d.get("experience", []) for b in e.get("bullets", [])]
    clamp = lambda x: max(0, min(100, int(round(x))))
    ats = ats_score(d, jd)

    must = _as_list((analysis or {}).get("must_have"))
    role_family = round(100 * sum(1 for k in must if _contains(txt, k.lower())) / len(must)) \
        if must else ats

    # Seniority alignment: penalize BOTH under-leveling and inflation (item 12).
    seniority = 100
    titles_blob = json.dumps([e.get("title", "") for e in d.get("experience", [])]).lower()
    if any(b in titles_blob for b in ("junior", "associate", "intern", "entry", "trainee", "graduate")):
        seniority = 60
    seniority -= 12 * len(_over_leveled_titles(d))

    cred = 100
    cred -= 8 * sum(1 for p in GENERIC_SUMMARY_PHRASES if p in txt)
    cred -= 6 * sum(1 for b in bullets for p in VAGUE_PHRASES
                    if re.search(r"\b" + re.escape(p) + r"\b", b.lower()))
    cred -= 10 * (len(_umbrella_skill_lines(d)) + len(_keyword_dump_lines(d)))
    cred -= 5 * len(_overused_keywords(d))          # forced keyword repetition
    cred -= 4 * len(_robotic_openers(bullets))      # robotic verb openers
    if _micro_category_lines(d):
        cred -= 10
    if _metric_monotony(bullets):
        cred -= 12
    # Over-quantification (every bullet a number) reads generated, per the feedback.
    if bullets and sum(1 for b in bullets if re.search(r"\d", b)) / len(bullets) > 0.9:
        cred -= 8

    # Interview defensibility: penalize JD-copying, tool-heavy bullets, over-tailored old roles.
    demo_required = [k.lower() for k in _as_list((analysis or {}).get("must_have"))]
    defens = 100 - 12 * len(_verbatim_overlaps(bullets, jd))
    defens -= 6 * len(_tool_heavy_bullets(bullets))
    defens -= 8 * len(_over_tailored_old_roles(d.get("experience", []), demo_required))

    fmt = 100 - 10 * len(_encoding_artifacts(d))
    fmt -= 6 * len(_ats_token_lines(d))            # fake hyphen/underscore tokens
    if "—" in json.dumps(d):
        fmt -= 15
    if _estimate_lines(d) > PAGE_LINES * 3:
        fmt -= 15

    impact_words = ("revenue", "cost", "latency", "throughput", "adoption", "retention",
                    "sla", "uptime", "accuracy", "conversion", "efficiency", "saved",
                    "reduced", "increased", "improved")
    impact_cov = sum(1 for b in bullets if any(w in b.lower() for w in impact_words))
    hm = round(100 * impact_cov / len(bullets)) if bullets else 0
    signals = _as_list((analysis or {}).get("success_signals"))
    if signals:  # matching the JD's success signals is a BOOST, never a drag
        sig_cov = sum(1 for s in signals if _contains(txt, s.lower())) / len(signals)
        hm = round(min(100, hm + 25 * sig_cov))

    return {
        "ats_match": clamp(ats),
        "role_family_alignment": clamp(role_family),
        "seniority_alignment": clamp(seniority),
        "recruiter_credibility": clamp(cred),
        "interview_defensibility": clamp(defens),
        "formatting_quality": clamp(fmt),
        "hiring_manager_fit": clamp(hm),
    }


# Length control (item 12): keep the document within a target page budget.
PAGE_LINES = 60  # approx body lines per page at 7.5pt with the current spacing


def _estimate_lines(d):
    import math
    cpl = 120  # approx characters per wrapped line across the 6.5in text column
    wrapped = lambda s: max(1, math.ceil(len(s) / cpl))
    lines = 4  # name + contact + spacing
    lines += 1 + sum(wrapped(s) for s in d.get("summary", []))
    lines += 1 + sum(wrapped(s) for s in d.get("skills", []))
    lines += 1  # experience heading
    for e in d.get("experience", []):
        lines += 2 + sum(wrapped(b) for b in e.get("bullets", []))
    if d.get("projects"):
        lines += 1
        for p in d["projects"]:
            lines += 1 + sum(wrapped(b) for b in p.get("bullets", []))
    lines += 1 + len(d.get("education", []))
    if d.get("certifications"):
        lines += 1 + len(d.get("certifications", []))
    return lines


def _trim_to_pages(d, max_pages=3):
    """If over the page budget, drop the least-relevant (last) bullets from the OLDEST
    roles first, never below 4 per role and never touching the most recent role."""
    budget = PAGE_LINES * max_pages
    for e in reversed(d.get("experience", [])[1:]):
        while _estimate_lines(d) > budget and len(e.get("bullets", [])) > 4:
            e["bullets"].pop()
    if _estimate_lines(d) > budget:
        for p in d.get("projects", []):
            while _estimate_lines(d) > budget and len(p.get("bullets", [])) > 1:
                p["bullets"].pop()
    return d


# Deterministic text cleanup (item 11): normalize quotes/dashes/encoding, fix the
# possessive corruption ("Lowe20s" -> "Lowe's"), collapse spacing, dedupe bullets.
_QUOTE_MAP = {
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "“": '"', "”": '"', "„": '"',
    "–": "-", "—": "-", "―": "-", "−": "-",
    "…": "...", " ": " ", "​": "",
    "\x91": "'", "\x92": "'", "\x93": '"', "\x94": '"',
    "\x95": "-", "\x96": "-", "\x97": "-", "\x85": "...",
}


def _normalize_text(s):
    if not isinstance(s, str):
        return s
    import re
    for bad, good in _QUOTE_MAP.items():
        s = s.replace(bad, good)
    s = re.sub(r"\b([A-Za-z]{2,})20s\b", r"\1's", s)  # Lowe20s -> Lowe's, Lee20s -> Lee's
    # Human-language cleanup: split fake ATS tokens (snake_case, 3+ part hyphen chains).
    s = re.sub(r"(?<=[A-Za-z])_(?=[A-Za-z])", " ", s)  # model_registry -> model registry
    s = re.sub(r"\b[A-Za-z]{2,}(?:-[A-Za-z]{2,}){2,}\b",
               lambda m: m.group(0) if m.group(0).lower() in _LEGIT_HYPHENS
               else m.group(0).replace("-", " "), s)   # agentic-coding-workflows -> spaces
    s = re.sub(r"\b([A-Za-z]{2,})-([A-Za-z]{2,})\b",   # LLM-integration -> LLM integration
               lambda m: (m.group(1) + " " + m.group(2))
               if (m.group(0).lower() not in _LEGIT_HYPHENS
                   and m.group(2).lower() in _ATS_NOUN_SECONDS)
               else m.group(0), s)
    s = s.replace("�", "")
    s = "".join(c for c in s if c == "\n" or ord(c) >= 32)
    s = re.sub(r"[ \t]{2,}", " ", s)            # collapse doubled spaces
    s = re.sub(r" +([,.;:%])", r"\1", s)        # no space before punctuation
    return s.strip()


def _normalize_deep(v):
    if isinstance(v, str):
        return _normalize_text(v)
    if isinstance(v, list):
        return [_normalize_deep(x) for x in v]
    if isinstance(v, dict):
        return {k: _normalize_deep(x) for k, x in v.items()}
    return v


def _dedupe_bullets(items):
    import re
    seen, out = set(), []
    for it in items:
        key = re.sub(r"\W+", "", (it or "").lower())
        if key and key not in seen:
            seen.add(key)
            out.append(it)
    return out


def _normalize_resume(d):
    out = _normalize_deep(d)
    for e in out.get("experience", []):
        e["bullets"] = _dedupe_bullets(e.get("bullets", []))
    for p in out.get("projects", []):
        p["bullets"] = _dedupe_bullets(p.get("bullets", []))
    return out


def data_to_text(d):
    """Plain-text preview mirroring the resume structure."""
    out = [d.get("name", ""), d.get("contact", ""), "", "PROFESSIONAL SUMMARY"]
    out += d.get("summary", [])
    out += ["", "TECHNICAL SKILLS"] + d.get("skills", [])
    out += ["", "PROFESSIONAL EXPERIENCE"]
    for e in d.get("experience", []):
        out += ["", f"{e.get('title','')}    {e.get('dates','')}", e.get("company_location", "")]
        out += [f"- {b}" for b in e.get("bullets", [])]
    if d.get("projects"):
        out += ["", "PROJECT HIGHLIGHTS"]
        for p in d["projects"]:
            out += ["", p.get("name", "")] + [f"- {b}" for b in p.get("bullets", [])]
    out += ["", "EDUCATION"] + d.get("education", [])
    if d.get("certifications"):
        out += ["", "Certifications"] + d.get("certifications", [])
    if d.get("analysis"):
        out += ["", "=" * 60, "ATS ANALYSIS", "=" * 60, d["analysis"]]
    return "\n".join(out)


# ---------------------------------------------------------------------------
# DOCX rendering — reproduces the original styling (Century Gothic, sizes, etc.)
# ---------------------------------------------------------------------------
def _set_font(run, size, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold


def _add_bottom_border(p):
    """Draw a horizontal divider line under a paragraph (section heading)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_hyperlink(paragraph, url, text, size=10.0):
    r_id = paragraph.part.relate_to(url, HYPERLINK_REL, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    rpr.append(fonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(text.upper()), 7.5, bold=True)
    _add_bottom_border(p)
    return p


def _body(doc, text, justify=False, bold=False, size=7.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(p.add_run(text), size, bold=bold)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(p.add_run(text), 7.5)
    return p


def _xml_safe(s):
    """Normalize quotes/dashes/encoding artifacts and strip control chars for the DOCX."""
    return _normalize_text(s)


def _clean(obj):
    if isinstance(obj, str):
        return _xml_safe(obj)
    if isinstance(obj, list):
        return [_clean(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _clean(v) for k, v in obj.items()}
    return obj


def render_docx(d, out_path, contact=None):
    d = _clean(d)
    # Per-profile contact overrides; fall back to the default constants. The contact
    # may also ride along inside the data dict (set by generate flow) as "_contact".
    contact = contact or d.get("_contact") or {}
    c_email = contact.get("email", EMAIL)
    c_phone = contact.get("phone", PHONE)
    c_links = {
        "LinkedIn": contact.get("linkedin", LINKS["LinkedIn"]),
        "GitHub": contact.get("github", LINKS["GitHub"]),
        "Portfolio": contact.get("portfolio", LINKS["Portfolio"]),
    }
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(7.5)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    _set_font(name_p.add_run(contact.get("name") or d.get("name", "")), 14, bold=True)

    # Contact line with real hyperlinks: email | phone | LinkedIn | GitHub | Portfolio
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(cp.add_run(c_email), 10)
    _set_font(cp.add_run(f"  |  {c_phone}  |  "), 10)
    for i, (label, url) in enumerate(c_links.items()):
        if i:
            _set_font(cp.add_run("  |  "), 10)
        _add_hyperlink(cp, url, label, 10)

    _heading(doc, "Professional Summary")
    for line in d.get("summary", []):
        _bullet(doc, line)

    _heading(doc, "Technical Skills")
    for line in d.get("skills", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if ":" in line:
            cat, rest = line.split(":", 1)
            _set_font(p.add_run(cat + ":"), 7.5, bold=True)
            _set_font(p.add_run(rest), 7.5)
        else:
            _set_font(p.add_run(line), 7.5)

    _heading(doc, "Professional Experience")
    for e in d.get("experience", []):
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(3)
        head.paragraph_format.space_after = Pt(0)
        # title left, dates right via a tab stop
        from docx.enum.text import WD_TAB_ALIGNMENT
        head.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        r1 = head.add_run(e.get("title", "")); _set_font(r1, 7.5, bold=True)
        r1.add_tab()
        r2 = head.add_run(e.get("dates", "")); _set_font(r2, 7.5, bold=True)
        _body(doc, e.get("company_location", ""), bold=True)
        for b in e.get("bullets", []):
            _bullet(doc, b)

    if d.get("projects"):
        _heading(doc, "Project Highlights")
        for p in d["projects"]:
            _body(doc, p.get("name", ""), bold=True)
            for b in p.get("bullets", []):
                _bullet(doc, b)

    _heading(doc, "Education")
    for line in d.get("education", []):
        _body(doc, line)

    if d.get("certifications"):
        _heading(doc, "Certifications")
        for line in d.get("certifications", []):
            _body(doc, line, bold=True)

    doc.save(out_path)
    return out_path
